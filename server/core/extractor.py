from pypdf import PdfReader
from docx import Document
import pandas as pd
import json
import io

def extract_txt(file):
    return file.file.read().decode("utf-8")

def extract_pdf(file):
    reader = PdfReader(file.file)
    text = ""
    
    for page in reader.pages:
        text += (page.extract_text() or "") + "\n"
        
    return text

def extract_docx(file):
    doc = Document(file.file)
    
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
        
    return text

def extract_csv(file):
    df = pd.read_csv(file.file)
    return df.to_string()

def extract_json(file):
    data = json.loads(file.file.read().decode("utf-8"))
    return json.dumps(data, indent=2)

def extract_md(file):
    return file.file.read().decode("utf-8")

def extract_text(file):
    filename = file.filename.lower()

    if filename.endswith(".txt"):
        return extract_txt(file)
    elif filename.endswith(".pdf"):
        return extract_pdf(file)
    elif filename.endswith(".docx"):
        return extract_docx(file)
    elif filename.endswith(".csv"):
        return extract_csv(file)
    elif filename.endswith(".json"):
        return extract_json(file)
    elif filename.endswith(".md"):
        return extract_md(file)
    else:
        raise ValueError("Unsupported file type")


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    buf = io.BytesIO(content)
    name = filename.lower()

    if name.endswith(".txt") or name.endswith(".md"):
        return content.decode("utf-8")
    elif name.endswith(".pdf"):
        reader = PdfReader(buf)
        return "".join((page.extract_text() or "") + "\n" for page in reader.pages)
    elif name.endswith(".docx"):
        doc = Document(buf)
        return "\n".join(para.text for para in doc.paragraphs)
    elif name.endswith(".csv"):
        df = pd.read_csv(buf)
        return df.to_string()
    elif name.endswith(".json"):
        data = json.loads(content.decode("utf-8"))
        return json.dumps(data, indent=2)
    else:
        raise ValueError(f"Unsupported file type: {filename}")
