import os

from docx import Document

def export_to_docx(messages):
    document = Document()
    document.add_heading("AI Knowledge Assistant", level=1)

    for index, message in enumerate(messages, start=1):
        document.add_heading(
            f"Chat {index}",
            level=2
        )

        document.add_heading(
            "Question",
            level=3
        )

        document.add_paragraph(message["question"])

        document.add_heading(
            "Answer",
            level=3
        )

        document.add_paragraph(message["answer"])

        document.add_heading(
            "Citations",
            level=3
        )

        if message["citations"]:
            for citation in message["citations"]:
                document.add_paragraph(
                    citation,
                    style="List Bullet"
                )
        else:
            document.add_paragraph("No citations.")

    os.makedirs("exports", exist_ok=True)

    output_path = "exports/response.docx"
    document.save(output_path)

    return output_path
